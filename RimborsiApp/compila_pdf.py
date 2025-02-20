from itertools import zip_longest

import io
import os
import re
from PyPDF2 import PdfFileReader, PdfFileWriter
from PyPDF2.generic import BooleanObject, IndirectObject, NameObject
from django.contrib.auth.decorators import login_required
from django.db.models import Sum
from django.http import HttpResponseBadRequest
from django.shortcuts import get_object_or_404, redirect, render
from docx import Document
from docx.shared import Cm, Inches
from reportlab.pdfgen import canvas

from .forms import *
from .views import money_exchange, resoconto_data, firma
from PIL import Image

from django.http import Http404

from reportlab.pdfgen import canvas
from PyPDF2 import PdfFileWriter, PdfFileReader
import io

from django.core.exceptions import ValidationError

@login_required
def genera_pdf(request, id):
    if request.method == 'POST':
        moduli_missione = get_object_or_404(ModuliMissione, missione_id=id)
        moduli_missione_form = ModuliMissioneForm(request.POST, instance=moduli_missione)
        firme_form = FirmaChooseForm(request.POST, user_owner=request.user, instance=moduli_missione)
        dichiarazione_check_pers = None
        dichiarazione_check_std = None
        if moduli_missione_form.is_valid() and firme_form.is_valid():
            moduli_missione_form.save()
            # return redirect('profile')
            dichiarazione_check_pers = moduli_missione_form.cleaned_data['dichiarazione_check_pers']
            dichiarazione_check_std = moduli_missione_form.cleaned_data['dichiarazione_check_std']
            firme_form.save()
            #moduli_missione.save()
        else:
            km, indennita, totali = resoconto_data(missione=moduli_missione.missione)
            return render(request, 'Rimborsi/resoconto.html', {'missione': moduli_missione.missione,
                                                               'moduli_missione_form': moduli_missione_form,
                                                               'km': km,
                                                               'indennita': indennita,
                                                               'totali': totali,
                                                               # Firme
                                                               'firme_form': firme_form,
                                                               })

        # Firme
        firma_richiedente = moduli_missione.firma_richiedente
        firma_titolare = moduli_missione.firma_titolare

        compila_anticipo(request, id, firma_richiedente, firma_titolare)
        compila_parte_1(request, id, firma_richiedente, firma_titolare)
        compila_parte_2(request, id, firma_richiedente, firma_titolare)
        if request.user.profile.qualifica == 'DOTTORANDO':
            compila_autorizz_dottorandi(request, id)
        compila_atto_notorio(request, id, firma_richiedente, dichiarazione_check_std, dichiarazione_check_pers)

        try:
            genera_report_scontrini(request, id)
        except Exception as e:
            print(e)

        return redirect('RimborsiApp:resoconto', id)
    else:
        return HttpResponseBadRequest()


def add_new_pdf(pdf_writer, img_obj):

    # Get path
    try:
        filename = img_obj.path
    except:
        # No file associated with the current object
        return

    if filename.endswith('.pdf'):
        # It's a PDF file
        pdf_reader = PdfFileReader(filename)

        # Check whether the pdf id actually fine (it should be improved)
        try:
            reader = PdfFileReader(filename)
            print("Opening '{}', pages={}".format(filename, reader.getNumPages()))
            # Try to write it into an dummy ByteIO stream to check whether pdf is broken
            writer = PdfFileWriter()
            writer.addPage(reader.getPage(0))
            writer.write(io.BytesIO())
        except:
            print("Error reading '{}".format(filename))
            return

        for page_num in range(pdf_reader.getNumPages()):
            page = pdf_reader.getPage(page_num)
            pdf_writer.addPage(page)
    else:
        # Assume it's an image file
        try:
            image = Image.open(filename)
        except IOError:
            return

        if image.mode != 'RGB':
            image = image.convert('RGB')

        # Convert the image to a PDF page in memory
        image_pdf_bytes = io.BytesIO()
        image.save(image_pdf_bytes, format='PDF')
        image_pdf_bytes.seek(0)

        # Check whether the pdf id actually fine (it should be improved)
        try:
            reader = PdfFileReader(image_pdf_bytes)
            print("Opening '{}', pages={}".format(filename, reader.getNumPages()))
            # Try to write it into an dummy ByteIO stream to check whether pdf is broken
            writer = PdfFileWriter()
            writer.addPage(reader.getPage(0))
            writer.write(io.BytesIO())
        except:
            print("Error reading '{}".format(filename))
            return

        # Read the PDF page and add it to the writer
        image_pdf_reader = PdfFileReader(image_pdf_bytes)
        page = image_pdf_reader.getPage(0)
        pdf_writer.addPage(page)


def genera_report_scontrini(request, id):
    moduli_output_path = os.path.join(settings.MEDIA_ROOT, 'moduli')

    missione = Missione.objects.get(user=request.user, id=id)
    profile = Profile.objects.get(user=request.user)
    trasporti = Trasporto.objects.filter(missione=missione).order_by('data')
    pasti = Pasti.objects.filter(missione=missione).order_by('data')
    spese = SpesaMissione.objects.filter(missione=missione).order_by('spesa__data')

    pdf_writer = PdfFileWriter()
    prev_date = None
    for pasti_del_giorno in pasti:
        if pasti_del_giorno.data != prev_date:
            pass

        add_new_pdf(pdf_writer, pasti_del_giorno.img_scontrino1)
        add_new_pdf(pdf_writer, pasti_del_giorno.img_scontrino2)
        add_new_pdf(pdf_writer, pasti_del_giorno.img_scontrino3)

    for trasporto in trasporti:
        add_new_pdf(pdf_writer, trasporto.img_scontrino)

    for spesa in spese:
        add_new_pdf(pdf_writer, spesa.spesa.img_scontrino)

    # Write the combined PDF to a BytesIO object
    output_pdf_bytes = io.BytesIO()
    pdf_writer.write(output_pdf_bytes)
    output_pdf_bytes.seek(0)

    output_name_tmp = os.path.join(moduli_output_path, f'Missione_{missione.id}_prove_di_acquisto_tmp.pdf')
    outputStream = open(output_name_tmp, "wb")
    pdf_writer.write(outputStream)
    outputStream.close()

    # Salvo il pdf appena creato dentro a un FileField
    output_name = f'Missione_{missione.id}_prove_di_acquisto.pdf'
    outputStream = open(output_name_tmp, "rb")
    moduli_missione = ModuliMissione.objects.get(missione=missione)
    moduli_missione.prove_acquisto_file.save(output_name, outputStream)

    # Elimino il file temporaneo
    os.remove(output_name_tmp)


def compila_anticipo(request, id, firma_richiedente, firma_titolare):
    moduli_input_path = os.path.join(settings.STATIC_ROOT, 'RimborsiApp', 'moduli')
    moduli_output_path = os.path.join(settings.MEDIA_ROOT, 'moduli')

    input_file = os.path.join(moduli_input_path, 'ModuloAnticipo.docx')
    document = Document(input_file)
    missione = Missione.objects.get(user=request.user, id=id)
    date_richiesta = ModuliMissione.objects.get(missione=missione)
    profile = Profile.objects.get(user=request.user)
    # trasporto = Trasporto.objects.filter(missione=missione)
    # km_totali = trasporto.filter(mezzo='AUTO').aggregate(Sum('km'))['km__sum'] or 0
    #
    # # Remove trasporti that has 0 km
    # trasporto = trasporto.filter(costo__gt=0).order_by("data")

    class ParConfig:
        def __init__(self):
            self.config = []
            self.counter = 0

        def append(self, index, values, excludes=[]):
            self.config.append([index, values, excludes])

        def __getitem__(self, index):
            return self.config[index]

    config = ParConfig()
    config.append("Richiedente (Cognome Nome):", [f'{profile.user.last_name} {profile.user.first_name}'])
    config.append("Nato/a a", [f'{profile.luogo_nascita.name} il {profile.data_nascita.strftime("%d/%m/%Y")}'])
    config.append("Codice fiscale", [f'{profile.cf}'])
    config.append("Domicilio fiscale", [f'Via {profile.domicilio.via} {profile.domicilio.n}, {profile.domicilio.comune.name} ({profile.domicilio.provincia.codice_targa})'])
    config.append("Qualifica/Ruolo", [f'{profile.qualifica}'])
    config.append("Datore di lavoro", [f'{profile.datore_lavoro}'])
    config.append("Destinazione missione", [f'{missione.citta_destinazione} - {missione.stato_destinazione.nome}'])
    config.append("Motivazione missione", [f'{missione.motivazione}'])

    config.append("l’inizio della missione è prevista alle ore", [f'{missione.inizio_ora.strftime("%H:%M")} del {missione.inizio.strftime("%d/%m/%Y")}'])
    config.append("la durata della missione è prevista in giorni", [f'{missione.durata_gg.days}'])
    config.append("la spesa graverà sul progetto", [f'{missione.fondo}'])
    config.append("Data richiesta", [date_richiesta.anticipo.strftime('%d/%m/%Y')])

    # config.append('DICHIARA di aver compiuto la missione a', [
    #     f'{missione.citta_destinazione} - {missione.stato_destinazione.nome}',
    #     missione.inizio_ora.strftime('%H:%M'),
    #     missione.inizio.strftime('%d/%m/%Y'),
    #     missione.fine_ora.strftime('%H:%M'),
    #     missione.fine.strftime('%d/%m/%Y')])
    # config.append('DICHIARA di aver ricevuto', ['TODO'])
    # config.append('nel caso di utilizzo di mezzo proprio', [f'{km_totali}'])
    # config.append('che il costo del biglietto', ['cosa ci va?'])
    # config.append('che l’originale del fattura/ricevuta cumulativa', ['aaa', 'aaa', 'aaa'])
    # config.append('che il costo della fattura/ricevuta _', ['aa'])
    # config.append('che l’originale del fattura/ricevuta cumulativa, relativo a ___', ['aaa', 'aaa', 'aaa'])
    # config.append('che l’originale del fattura/ricevuta cumulativa, relativo a __________________________ ',
    #               ['aaa', 'aaa', 'aaa'])
    # config.append('che il costo della fattura/ricevuta __________________________ ', ['aa'])
    # config.append('Data richiesta', [date_richiesta.parte_2.strftime('%d/%m/%Y')])

    for k, values, excludes in config:
        str = ''
        for par in document.paragraphs:
            if k in par.text:
                for s1, s2 in zip_longest(re.sub('_+', '_', par.text).split('_'), values, fillvalue=''):
                    if s2 != '':
                        str += f'{s1}__{s2}__'
                    else:
                        str += f'{s1}'
                for r in par.runs:
                    if len(r.text) > 0:
                        r.text = ''
                par.add_run(text=str)
                break

    inserisci_firme(document, firma_richiedente, firma_titolare)

    output_name_tmp = os.path.join(moduli_output_path, f'Missione_{missione.id}_anticipo_tmp.docx')
    output_name = f'Missione_{missione.id}_anticipo.docx'
    document.save(os.path.join(moduli_output_path, output_name_tmp))
    # Salvo il docx appena creato dentro a un FileField
    moduli_missione = ModuliMissione.objects.get(missione=missione)
    outputStream = open(output_name_tmp, "rb")
    moduli_missione.anticipo_file.save(output_name, outputStream)

    # Elimino il file temporaneo
    os.remove(output_name_tmp)


def compila_parte_1(request, id, idR, idT):
    moduli_input_path = os.path.join(settings.STATIC_ROOT, 'RimborsiApp', 'moduli')
    moduli_output_path = os.path.join(settings.MEDIA_ROOT, 'moduli')

    input_file = os.path.join(moduli_input_path, 'ModuloMissione_Part1.pdf')
    coords_dict = {
        "struttura_appartenenza": [70, 748],
        "cognome": [170, 695],
        "nome": [370, 695],
        "luogo_nascita": [100, 681],
        "data_nascita": [400, 681],
        "cf": [133, 654],
        "domicilio": [150, 630],
        "prov": [481, 630],
        "qualifica": [140, 604],
        "datore": [347, 604],
        "destinazione": [165, 578],
        "motivazione": [165, 552],
        "inizio_ora": [140, 513],
        "inizio": [310, 513],
        "durata": [260, 486],
        "fondo": [218, 470],
        "mezzo": [133, 438],
        "motivazione_mezzi": [59, 408],
        "targa": [391, 360],
        "data_richiesta": [133, 180],
    }

    # Firma coords
    firma_coords_richiedente = [180, 150]
    firma_coords_titolare = [245, 125]

    missione = Missione.objects.get(user=request.user, id=id)
    date_richiesta = ModuliMissione.objects.get(missione=missione)
    profile = Profile.objects.get(user=request.user)
    trasporto = Trasporto.objects.filter(missione=missione)

    trasporto_set = set()
    for t in trasporto:
        trasporto_set.add(t.mezzo)

    targa = ""
    motivazione_auto = ""
    if "AUTO" in eval(missione.mezzi_previsti):
        try:
            targa = missione.automobile.targa
            motivazione_auto = ' + '.join(t.lower() for t in eval(missione.motivazione_automobile))
        except:
            print("Un untente ha selezionato 'AUTO' come mezzo per la missione, ma non hai compilato i dati relativi.")

    new_list = eval(missione.mezzi_previsti)
    if "A_ALT" in new_list:
        for n, elem in enumerate(new_list):
            if elem == "A_ALT":
                new_list[n] = f'AUTO ALTRUI (di proprietà di {missione.automobile_altrui or ""})'
    mezzo = ' + '.join(t for t in new_list)

    data_fine_rapporto_str = ''
    if profile.data_fine_rapporto is not None:
        data_fine_rapporto_str = ' (' + profile.data_fine_rapporto.strftime('%d/%m/%Y') + ')'
    qualifica_fine = profile.get_qualifica_display() + data_fine_rapporto_str

    interno_str = ''
    if profile.telefono is not None:
        interno_str = '   (interno: ' + str(profile.telefono) + ')'
    nome_interno = profile.user.first_name + interno_str

    if profile.straniero:
        luogo_nascita = profile.luogo_nascita_straniero
        domicilio = f'{profile.domicilio.via} {profile.domicilio.n}, {profile.domicilio.comune_straniero}'
        provincia = profile.domicilio.provincia_straniero
    else:
        luogo_nascita = profile.luogo_nascita.name
        domicilio = f'{profile.domicilio.via} {profile.domicilio.n}, {profile.domicilio.comune.name}'
        provincia = profile.domicilio.provincia.codice_targa

    value_dict = {
        "struttura_appartenenza": missione.struttura_fondi,
        "cognome": profile.user.last_name,
        "nome": nome_interno,
        "luogo_nascita": luogo_nascita,
        "data_nascita": profile.data_nascita.strftime('%d/%m/%Y'),
        "cf": profile.cf,
        "domicilio": domicilio,
        "prov": provincia,
        "qualifica": qualifica_fine,
        "datore": profile.datore_lavoro,
        "destinazione": missione.citta_destinazione + ", " + missione.stato_destinazione.nome,
        "motivazione": missione.motivazione,
        "inizio_ora": missione.inizio_ora.strftime('%H:%M'),
        "inizio": missione.inizio.strftime('%d/%m/%Y'),
        "durata": str(missione.durata_gg.days),
        "fondo": missione.fondo,
        # "mezzo": ' + '.join(t for t in trasporto_set),
        "mezzo": mezzo,
        "motivazione_mezzi": motivazione_auto,
        "targa": targa,
        "data_richiesta": date_richiesta.parte_1.strftime('%d/%m/%Y'),
    }

    buffer = io.BytesIO()
    can = canvas.Canvas(buffer)
    can.setFont("Times-Roman", 11)
    # Write values
    for k, v in coords_dict.items():
        if k == 'cf':
            can.setFont('Courier', 18.2)
            can.drawString(*v, value_dict[k])
            can.setFont("Times-Roman", 11)
        else:
            if value_dict[k] is None:
                value_dict[k] = ""
            can.drawString(*v, value_dict[k])

    inserisci_firme_pdf(can, idR, firma_coords_richiedente, idT, firma_coords_titolare)

    can.showPage()
    can.save()
    buffer.seek(0)
    new_pdf = PdfFileReader(buffer)

    # Leggo il file base
    input = PdfFileReader(open(input_file, "rb"))  # Base file
    page = input.getPage(0)
    # Faccio il merge delle modifiche con il file base
    page.mergePage(new_pdf.getPage(0))

    # Scrivo tutto in un file temporaneo
    output = PdfFileWriter()
    output.addPage(page)
    output_name_tmp = os.path.join(moduli_output_path, f'Missione_{missione.id}_parte_1_tmp.pdf')
    outputStream = open(output_name_tmp, "wb")
    output.write(outputStream)
    outputStream.close()

    # Salvo il pdf appena creato dentro a un FileField
    output_name = f'Missione_{missione.id}_parte_1.pdf'
    outputStream = open(output_name_tmp, "rb")
    moduli_missione = ModuliMissione.objects.get(missione=missione)
    moduli_missione.parte_1_file.save(output_name, outputStream)

    # Elimino il file temporaneo
    os.remove(output_name_tmp)


def compila_parte_2(request, id, firma_richiedente, firma_titolare):
    moduli_input_path = os.path.join(settings.STATIC_ROOT, 'RimborsiApp', 'moduli')
    moduli_output_path = os.path.join(settings.MEDIA_ROOT, 'moduli')

    input_file = os.path.join(moduli_input_path, 'ModuloMissione_Part2.docx')
    document = Document(input_file)
    missione = Missione.objects.get(user=request.user, id=id)
    date_richiesta = ModuliMissione.objects.get(missione=missione)
    profile = Profile.objects.get(user=request.user)
    trasporto = Trasporto.objects.filter(missione=missione)
    km_totali = trasporto.filter(mezzo='AUTO').aggregate(Sum('km'))['km__sum'] or 0

    # Remove trasporti that has 0 euros
    # trasporto = trasporto.filter(costo__gt=0).order_by("data")
    trasporto = trasporto.order_by("data")  # It is better to keep them and specify they are just for km refound

    class ParConfig:
        def __init__(self):
            self.config = []
            self.counter = 0

        def append(self, index, values, excludes=[]):
            self.config.append([index, values, excludes])

        def __getitem__(self, index):
            return self.config[index]

    config = ParConfig()
    sottoscritto_gender_friendly = 'Il sottoscritto' if profile.sesso == 'M' else 'La sottoscritta'
    config.append(sottoscritto_gender_friendly, [f'{profile.user.first_name} {profile.user.last_name}'])
    config.append('DICHIARA di aver compiuto la missione a', [
        f'{missione.citta_destinazione} - {missione.stato_destinazione.nome}',
        missione.inizio_ora.strftime('%H:%M'),
        missione.inizio.strftime('%d/%m/%Y'),
        missione.fine_ora.strftime('%H:%M'),
        missione.fine.strftime('%d/%m/%Y')])
    config.append('DICHIARA di aver ricevuto', [f'{missione.anticipo}'])
    config.append('nel caso di utilizzo di mezzo proprio', [f'{km_totali}'])
    # config.append('che il costo del biglietto', ['cosa ci va?'])
    # config.append('che l’originale del fattura/ricevuta cumulativa', ['aaa', 'aaa', 'aaa'])
    # config.append('che il costo della fattura/ricevuta _', ['aa'])
    # config.append('che l’originale del fattura/ricevuta cumulativa, relativo a ___', ['aaa', 'aaa', 'aaa'])
    # config.append('che l’originale del fattura/ricevuta cumulativa, relativo a __________________________ ',
    #               ['aaa', 'aaa', 'aaa'])
    # config.append('che il costo della fattura/ricevuta __________________________ ', ['aa'])
    config.append('Data richiesta', [date_richiesta.parte_2.strftime('%d/%m/%Y')])

    for k, values, excludes in config:
        str = ''
        for par in document.paragraphs:
            if k in par.text:
                for s1, s2 in zip_longest(re.sub('_+', '_', par.text).split('_'), values, fillvalue=''):
                    if s2 != '':
                        str += f'{s1}__{s2}__'
                    else:
                        str += f'{s1}'
                for r in par.runs:
                    if len(r.text) > 0:
                        r.text = ''
                par.add_run(text=str)
                break

    # Tabella `VIAGGIO E TRASPORTO`
    table = document.tables[0]

    while len(table.rows) <= len(trasporto):
        row = table.add_row()

    for i, t in enumerate(trasporto, start=1):
        costo_str = f'{t.costo:.2f} {t.valuta}'
        if t.valuta != 'EUR':
            costo_in_euro = money_exchange(t.data, t.valuta, t.costo)
            costo_str += f' ({costo_in_euro:.2f} EUR)'

        table.cell(i, 0).text = t.data.strftime('%d/%m/%Y')
        table.cell(i, 1).text = f'da {t.da or ""}'
        table.cell(i, 2).text = f'a {t.a or ""}'
        table.cell(i, 3).text = t.mezzo
        if t.costo == 0:
            t.tipo_costo = t.tipo_costo or ''
            t.tipo_costo += ' + ' if t.tipo_costo != '' else ''
            t.tipo_costo += f'Rimborso km {t.km}'
        table.cell(i, 4).text = t.tipo_costo or ''
        table.cell(i, 5).text = costo_str
        table.rows[i].height = Cm(0.61)

    # Recupero delle altre spese dal database
    pernottamenti = Spesa.objects.filter(spesamissione__missione=missione, spesamissione__tipo='PERNOTTAMENTO')
    pasti = Pasti.objects.filter(missione=missione)
    convegni = Spesa.objects.filter(spesamissione__missione=missione, spesamissione__tipo='CONVEGNO')
    altre_spese = Spesa.objects.filter(spesamissione__missione=missione, spesamissione__tipo='ALTRO')

    spese_dict = {
        'pernottamento': pernottamenti,
        'pasto': pasti,
        'convegno': convegni,
        'altro': altre_spese,
    }

    ##################################
    # BEGIN: Old json-based version
    ##################################
    # db_dict = {
    #     'pernottamento': [],
    #     'scontrino': [],  # pasti
    #     'convegno': [],
    #     'altrespese': [],
    # }

    # Load the default values for each field in db_dict
    # for k, _ in db_dict.items():
    #     db_dict[k] = load_json(missione, k)
    #
    # # TODO Scontrino dovrebbe essere ordinato per data. Potrebbe essere un pacco.
    # r_scontrino = []
    # for row in db_dict['scontrino']:
    #     if row['s1'] is not None:
    #         r_scontrino.append({'data': row['data'], 's1': row['s1'], 'd1': row['d1'], 'v1': row['v1']})
    #     if row['s2'] is not None:
    #         r_scontrino.append({'data': row['data'], 's1': row['s2'], 'd1': row['d2'], 'v1': row['v2']})
    #     if row['s3'] is not None:
    #         r_scontrino.append({'data': row['data'], 's1': row['s3'], 'd1': row['d3'], 'v1': row['v3']})
    # db_dict['scontrino'] = r_scontrino


    # # Fill all the remaining tables
    # for index, (key, value) in enumerate(db_dict.items(), start=1):
    #     table = document.tables[index]
    #
    #     while len(table.rows) <= len(value):
    #         row = table.add_row()
    #         # row.height = Cm(0.61)
    #
    #     for i, t in enumerate(value, start=1):
    #         s1 = float(t['s1'])
    #         valuta = t['v1']
    #         data = t['data']  #.strftime('%Y-%m-%d')
    #         costo_str = f'{s1:.2f} {valuta}'
    #         if valuta != 'EUR':
    #             costo_in_euro = money_exchange(data, valuta, s1)
    #             costo_str += f' ({costo_in_euro:.2f} EUR)'
    #
    #         table.cell(i, 0).text = t['data'].strftime('%d/%m/%Y')
    #         table.cell(i, 1).text = t['d1'] if t['d1'] is not None else ''
    #         table.cell(i, 2).text = costo_str
    #         table.rows[i].height = Cm(0.61)
    ##################################
    # END: Old json-based version
    ##################################

    # Fill all the remaining tables
    for index, (key, queryset) in enumerate(spese_dict.items(), start=1):
        table = document.tables[index]

        total_rows = queryset.count()
        if key == 'pasto':
            total_rows = sum(1 for spesa in queryset for j in range(1, 4) if getattr(spesa, f'importo{j}'))

        # while len(table.rows) <= total_rows:
        #     row = table.add_row()

        row_index = 1
        for spesa in queryset:
            if key == 'pasto':
                for j in range(1, 4):
                    importo = getattr(spesa, f'importo{j}')
                    if importo:
                        valuta = getattr(spesa, f'valuta{j}')
                        descrizione = getattr(spesa, f'descrizione{j}')
                        data = spesa.data
                        costo_str = f'{importo:.2f} {valuta}'
                        if valuta != 'EUR':
                            costo_in_euro = money_exchange(data, valuta, importo)
                            costo_str += f' ({costo_in_euro:.2f} EUR)'

                        if row_index >= len(table.rows):
                            table.add_row()

                        table.cell(row_index, 0).text = data.strftime('%d/%m/%Y')
                        table.cell(row_index, 1).text = descrizione if descrizione else ''
                        table.cell(row_index, 2).text = costo_str
                        table.rows[row_index].height = Cm(0.61)
                        row_index += 1
            else:
                importo = spesa.importo
                valuta = spesa.valuta
                descrizione = spesa.descrizione
                data = spesa.data
                costo_str = f'{importo:.2f} {valuta}'
                if valuta != 'EUR':
                    costo_in_euro = money_exchange(data, valuta, importo)
                    costo_str += f' ({costo_in_euro:.2f} EUR)'

                if row_index >= len(table.rows):
                    table.add_row()

                table.cell(row_index, 0).text = data.strftime('%d/%m/%Y')
                table.cell(row_index, 1).text = descrizione if descrizione else ''
                table.cell(row_index, 2).text = costo_str
                table.rows[row_index].height = Cm(0.61)
                row_index += 1

    inserisci_firme(document, firma_richiedente, firma_titolare)       # sezione aggiunta firma richiedente e titolare

    output_name_tmp = os.path.join(moduli_output_path, f'Missione_{missione.id}_parte_2_tmp.docx')
    output_name = f'Missione_{missione.id}_parte_2.docx'
    document.save(os.path.join(moduli_output_path, output_name_tmp))
    # Salvo il docx appena creato dentro a un FileField
    moduli_missione = ModuliMissione.objects.get(missione=missione)
    outputStream = open(output_name_tmp, "rb")
    moduli_missione.parte_2_file.save(output_name, outputStream)

    # Elimino il file temporaneo
    os.remove(output_name_tmp)


def compila_autorizz_dottorandi(request, id):
    moduli_input_path = os.path.join(settings.STATIC_ROOT, 'RimborsiApp', 'moduli')
    moduli_output_path = os.path.join(settings.MEDIA_ROOT, 'moduli')

    input_file = os.path.join(moduli_input_path, 'AutorizzazioneDottorandi.pdf')
    coords_dict = {
        'data_richiesta': [120, 728],
        'tutor': [165, 577],
        'nomecognome': [410, 577],
        'anno_dottorato': [160, 550],
        'scuola_dottorato': [150, 522],
        'destinazione': [350, 438],
        'inizio': [100, 410],
        'fine': [250, 410],
        'motivazione': [62, 383],

    }
    missione = Missione.objects.get(user=request.user, id=id)
    date_richiesta = ModuliMissione.objects.get(missione=missione)
    profile = Profile.objects.get(user=request.user)

    value_dict = {
        'data_richiesta': date_richiesta.dottorandi.strftime('%d/%m/%Y'),
        'tutor': profile.tutor,
        'nomecognome': f'{profile.user.first_name} {profile.user.last_name}',
        'anno_dottorato': f'{profile.anno_dottorato if profile.anno_dottorato is not None else ""}',
        'scuola_dottorato': profile.scuola_dottorato,
        'destinazione': f'{missione.citta_destinazione} - {missione.stato_destinazione.nome}',
        'inizio': missione.inizio.strftime('%d/%m/%Y'),
        'fine': missione.fine.strftime('%d/%m/%Y'),
        'motivazione': missione.motivazione,
    }

    buffer = io.BytesIO()
    can = canvas.Canvas(buffer)
    can.setFont("Times-Roman", 11)
    # Write values
    for k, v in coords_dict.items():
        if k == 'cf':
            can.setFont('Courier', 18.2)
            can.drawString(*v, value_dict[k])
            can.setFont("Times-Roman", 11)
        else:
            if value_dict[k] is None:
                value_dict[k] = ""
            can.drawString(*v, value_dict[k])

    can.showPage()
    can.save()
    buffer.seek(0)
    new_pdf = PdfFileReader(buffer)

    # Leggo il file base
    input = PdfFileReader(open(input_file, "rb"))  # Base file
    page = input.getPage(0)
    # Faccio il merge delle modifiche con il file base
    page.mergePage(new_pdf.getPage(0))

    # Scrivo tutto in un file temporaneo
    output = PdfFileWriter()
    output.addPage(page)
    output_name_tmp = os.path.join(moduli_output_path, f'Missione_{missione.id}_autoriz_dottorandi_tmp.pdf')
    outputStream = open(output_name_tmp, "wb")
    output.write(outputStream)
    outputStream.close()

    # Salvo il pdf appena creato dentro a un FileField
    output_name = f'Missione_{missione.id}_autoriz_dottorandi.pdf'
    outputStream = open(output_name_tmp, "rb")
    moduli_missione = ModuliMissione.objects.get(missione=missione)
    moduli_missione.dottorandi_file.save(output_name, outputStream)

    # Elimino il file temporaneo
    os.remove(output_name_tmp)


def set_need_appearances_writer(writer):
    try:
        catalog = writer._root_object
        # get the AcroForm tree and add "/NeedAppearances attribute
        if "/AcroForm" not in catalog:
            writer._root_object.update({NameObject("/AcroForm"): IndirectObject(len(writer._objects), 0, writer)})
        need_appearances = NameObject("/NeedAppearances")
        writer._root_object["/AcroForm"][need_appearances] = BooleanObject(True)
    except Exception as e:
        print('set_need_appearances_writer() catch : ', repr(e))

    return writer


def compila_atto_notorio(request, id, firma_richiedente, dichiarazione_check_std=False, dichiarazione_check_pers=False):
    moduli_input_path = os.path.join(settings.STATIC_ROOT, 'RimborsiApp', 'moduli')
    moduli_output_path = os.path.join(settings.MEDIA_ROOT, 'moduli')
    missione = Missione.objects.get(user=request.user, id=id)
    input_file = os.path.join(moduli_input_path, 'dichiarazione_atto_notorieta.pdf')
    modulo_missione = ModuliMissione.objects.get(missione=missione)

    # open the pdf
    input_stream = open(input_file, "rb")
    pdf_reader = PdfFileReader(input_stream, strict=False)
    if "/AcroForm" in pdf_reader.trailer["/Root"]:
        pdf_reader.trailer["/Root"]["/AcroForm"].update({NameObject("/NeedAppearances"): BooleanObject(True)})

    pdf_writer = PdfFileWriter()
    set_need_appearances_writer(pdf_writer)
    if "/AcroForm" in pdf_writer._root_object:
        # Acro form is form field, set needs appearances to fix printing issues
        pdf_writer._root_object["/AcroForm"].update({NameObject("/NeedAppearances"): BooleanObject(True)})

    dichiarazione = ''
    if dichiarazione_check_std:
        dichiarazione += f'Di essersi recato a {missione.citta_destinazione} - {missione.stato_destinazione.nome} dal' \
                         f' {missione.inizio.strftime("%d/%m/%Y")} al {missione.fine.strftime("%d/%m/%Y")}' \
                         f' per {missione.motivazione}.'
    if dichiarazione_check_pers:
        dichiarazione += f' {modulo_missione.atto_notorio_dichiarazione}'

    if missione.user.profile.straniero:
        luogo_nascita = missione.user.profile.luogo_nascita_straniero
        provincia_nascita = ''
        residenza_comune = missione.user.profile.residenza.comune_straniero
        residenza_provincia = missione.user.profile.residenza.provincia_straniero
        domicilio_comune = missione.user.profile.domicilio.comune_straniero
        domicilio_provincia = missione.user.profile.domicilio.provincia_straniero
    else:
        luogo_nascita = missione.user.profile.luogo_nascita.name
        provincia_nascita = missione.user.profile.luogo_nascita.provincia.codice_targa
        residenza_comune = missione.user.profile.residenza.comune.name
        residenza_provincia = missione.user.profile.residenza.provincia.codice_targa
        domicilio_comune = missione.user.profile.domicilio.comune.name
        domicilio_provincia = missione.user.profile.domicilio.provincia.codice_targa

    data_dict = {
        '1': missione.user.last_name,
        '2': missione.user.first_name,
        # '3': missione.user.profile.luogo_nascita.name,
        '3': luogo_nascita,
        # '4': missione.user.profile.luogo_nascita.provincia.codice_targa,
        '4': provincia_nascita,
        '5': missione.user.profile.data_nascita.strftime('%d/%m/%Y'),
        # '6': missione.user.profile.residenza.comune.name,
        '6': residenza_comune,
        # '7': missione.user.profile.residenza.provincia.codice_targa,
        '7': residenza_provincia,
        '8': missione.user.profile.residenza.via,
        '9': missione.user.profile.residenza.n,
        # '10': missione.user.profile.domicilio.comune.name,
        '10': domicilio_comune,
        # '11': missione.user.profile.domicilio.provincia.codice_targa,
        '11': domicilio_provincia,
        '12': missione.user.profile.domicilio.via,
        '13': missione.user.profile.domicilio.n,
        '14': dichiarazione,
        '20': f'Modena, {modulo_missione.atto_notorio.strftime("%d/%m/%Y")}',
    }

    pdf_writer.addPage(pdf_reader.getPage(0))
    page = pdf_writer.getPage(0)
    pdf_writer.updatePageFormFieldValues(page, data_dict)

    # Sezione firma
    if firma_richiedente:
        try:
            firma_richiedente_img_path = firma_richiedente.img_firma.path
            firma_pdf = crea_firma_pdf(firma_richiedente_img_path)
            # Unisci la pagina della firma con il PDF originale
            page = pdf_writer.getPage(0)
            page.mergePage(firma_pdf.getPage(0))  # Sovrapponi la pagina della firma
        except Http404:
            print(f"Firma del richiedente {firma_richiedente} non trovata.")
            # TODO error handling should be rethought in the following functions

    # Disable the fillable fields
    # for j in range(0, len(page['/Annots'])):
    #     writer_annot = page['/Annots'][j].getObject()
    #     writer_annot.update({NameObject("/Ff"): NumberObject(1)})

    output_name_tmp = os.path.join(moduli_output_path, f'Missione_{missione.id}_atto_notorio_tmp.pdf')
    outputStream = open(output_name_tmp, "wb")
    pdf_writer.write(outputStream)
    outputStream.close()

    # Salvo il pdf appena creato dentro a un FileField
    output_name = f'Missione_{missione.id}_atto_notorio.pdf'
    outputStream = open(output_name_tmp, "rb")
    modulo_missione.atto_notorio_file.save(output_name, outputStream)

    # Elimino il file temporaneo
    os.remove(output_name_tmp)

# TODO error handling should be rethought in the following functions
def inserisci_firme(document, firma_richiedente, firma_titolare):

    if firma_richiedente:
        try:
            firma_richiedente_img_path = firma_richiedente.img_firma.path
            inserisci_firma( firma_richiedente_img_path, document , "Firma del richiedente") # Inserisce la firma del richiedente nel documento
            if os.path.exists(firma_richiedente_img_path):                                       # Verifica se il file esiste
                inserisci_firma(firma_richiedente_img_path, document, "Firma del richiedente")
            else:
                print(f"Firma del richiedente non trovata nel percorso: {firma_richiedente_img_path}")

        except Http404:
            print(f"Firma del richiedente {firma_richiedente} non trovata.")

    if firma_titolare:
        try:
            firma_titolare_img_path = firma_titolare.img_firma.path                                             # Firma_Shared non ha img-firma !
            inserisci_firma( firma_titolare_img_path, document, "Firma del titolare dei fondi/progetto" )         # Inserisce la firma del titolare nel documento
            if os.path.exists(firma_titolare_img_path):                                                               # Verifica se il file esiste
                inserisci_firma(firma_titolare_img_path, document, "Firma del titolare dei fondi/progetto")
            else:
                print(f"Firma del titolare non trovata nel percorso: {firma_titolare_img_path}")

        except Http404:
            print(f"Firma del titolare {firma_titolare} non trovata.")


def inserisci_firma(firma_img_path, document, str):
    for par in document.paragraphs:
        if str in par.text:

            match = re.search(r"_{3,}", par.text)  # Cerca dove iniziano 3 o più sottolineature di fila
            if match:
                start_text = par.text[:match.start()] # Testo prima delle linee
                end_text = par.text[match.end():] # Testo dopo le linee (se esiste)

                for r in par.runs: # Cancella il testo attuale del paragrafo
                    r.text = ""

                # ricostruzione del paragrafo con :
                par.add_run(start_text)     #  il testo prima delle linee
                par.add_run("__")           #  una parte delle sottolineature prima della firma
                par.add_run().add_picture(firma_img_path, width=Inches(0.8))        #  l'immagine della firma
                par.add_run("__")           #  la parte finale delle sottolineature dopo la firma
                par.add_run(end_text)       # il continuo del testo
            break


def inserisci_firme_pdf(can, firma_richiedente, coordR, firma_titolare, coordT):
    if firma_richiedente:
        try:
            firma_richiedente_img_path = firma_richiedente.img_firma.path
            can.drawImage(firma_richiedente_img_path, coordR[0], coordR[1], width=75, height=20,mask='auto')
        except Http404:
            print(f"Firma del richiedente {firma_richiedente} non trovata.")

    if firma_titolare:
        try:
            firma_titolare_img_path = firma_titolare.img_firma.path
            can.drawImage(firma_titolare_img_path, coordT[0], coordT[1], width=75, height=20, mask='auto')
        except Http404:
            print(f"Firma del titolare {firma_titolare} non trovata.")


def crea_firma_pdf(firma_path, firma_coords=[380, 150], firma_size=(100, 100)):
    """
    Crea un PDF temporaneo con l'immagine della firma.

    :return: PdfFileReader oggetto con il PDF contenente la firma
    """
    buffer = io.BytesIO()
    can = canvas.Canvas(buffer)
    can.drawImage(firma_path, firma_coords[0], firma_coords[1], width=firma_size[0], height=firma_size[1], mask='auto')
    can.showPage()
    can.save()
    buffer.seek(0)
    return PdfFileReader(buffer)